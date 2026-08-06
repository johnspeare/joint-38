"""Shared constants and Markdown preprocessing for the docx/PDF build pipeline.

Both build_sog_docx.py and build_sog_pdf.py start from the same Markdown
source and need the same fixups before handing it to Pandoc:

1. Swap the chain-of-command SVG for its PNG counterpart. Word/LibreOffice's
   SVG support can't render mermaid-cli's <foreignObject> node labels any
   more reliably than WeasyPrint could (see handoff.md §9) — same
   workaround, same PNG asset, new consumer.
2. Cap that image's height. Pandoc places images at native pixel size with
   no default max-height, so the chain-of-command PNG (546x1986px) gets
   inserted at ~21 inches tall and spills across several pages. A single
   `{height=...}` attribute (needs the `attributes` Pandoc extension) fixes
   it, matching the `max-height: 6.5in` rule already in the old WeasyPrint
   CSS.
3. Insert the title-page logo right after the H1. This has to be shared,
   not PDF-only: restyle_title_page() (below) assumes the paragraph right
   after the H1 is the logo image, and silently mangles whatever paragraph
   *is* there if that assumption is wrong — this used to be PDF-only,
   which meant the docx build had no logo there at all, so
   restyle_title_page() was centering/huge-spacing the "Introduction"
   heading and force-squashing the Chain of Command diagram (the actual
   first image in that build) into a 3.25"x3.25" square.
"""

from __future__ import annotations

import re
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

DOC_BASENAME = "Joint Fire 3&8 Standard Operating Guide DRAFT"
TITLE_PAGE_LINE_1 = "Joint Fire Protection 3 & 8"
TITLE_PAGE_LINE_2 = "Standard Operating Guidelines"
TITLE_PAGE_REVISION = "Revised Summer 2026"
TITLE_LOGO_SIZE = Inches(3.25)

REPO_ROOT = Path(__file__).resolve().parent.parent
ASSETS_DIR = REPO_ROOT / "FD-SOGs-assets"
SOURCE_MD = REPO_ROOT / "sog-1st-pass.md"
REFERENCE_DOCX = Path(__file__).resolve().parent / "reference.docx"

_CHAIN_OF_COMMAND_RE = re.compile(
    r"!\[([^\]]*)\]\(FD-SOGs-assets/chain-of-command\.svg\)"
)
_TITLE_LINE_RE = re.compile(r"^# .+\n", re.MULTILINE)
TITLE_LOGO_IMAGE = (ASSETS_DIR / "38-logo-title.png").as_posix()
TITLE_LOGO_WIDTH = "2.5in"
_TITLE_AND_LOGO_RE = re.compile(
    r"^# .+\n\n!\[\]\(" + re.escape(TITLE_LOGO_IMAGE) + r"\)\{width=" + re.escape(TITLE_LOGO_WIDTH) + r"\}\n",
    re.MULTILINE,
)


def preprocess_markdown(md_text: str) -> str:
    """Apply the docx/PDF-only fixups described above to Markdown source text."""
    png_path = (ASSETS_DIR / "chain-of-command.png").as_posix()

    def _swap(m: re.Match[str]) -> str:
        alt = m.group(1)
        return f"![{alt}]({png_path}){{height=6in}}"

    md_text = _CHAIN_OF_COMMAND_RE.sub(_swap, md_text)

    m = _TITLE_LINE_RE.match(md_text)
    if not m:
        raise ValueError("expected the Markdown source to start with a single H1 title line")
    title_line = m.group(0)
    rest = md_text[m.end():]
    md_text = f"{title_line}\n![]({TITLE_LOGO_IMAGE}){{width={TITLE_LOGO_WIDTH}}}\n" + rest

    return md_text


def load_preprocessed_source() -> str:
    md_text = SOURCE_MD.read_text(encoding="utf-8")
    return preprocess_markdown(md_text)


def split_after_title_logo(md_text: str) -> tuple[str, str]:
    """Split already-preprocessed Markdown right after the title+logo block.

    For build_sog_pdf.py, which needs to insert its page-break/TOC scaffold
    right after the (already shared-inserted) title and logo, without
    re-inserting a second logo.
    """
    m = _TITLE_AND_LOGO_RE.match(md_text)
    if not m:
        raise ValueError("expected preprocessed Markdown to start with the title heading + logo image block")
    return m.group(0), md_text[m.end():]


def restyle_title_page(docx_path: Path) -> None:
    """Post-process a Pandoc-generated docx's title page (shared by both builds).

    Pandoc has no concept of "the cover page" — the H1 and the logo that
    follows it are just the first two paragraphs of the body. This finds
    them structurally (the H1 is unique in this document) and rewrites:
    the title to a short, centered, two-line form; the logo, centered and
    enlarged; and adds a centered "Revised ..." line pushed down near the
    bottom of the page via paragraph spacing (no true frame/anchor needed
    since it's the only content on the page other than the title/logo).
    """
    doc = Document(str(docx_path))

    paragraphs = doc.paragraphs
    h1 = next(p for p in paragraphs if p.style.name == "Heading 1")
    h1_index = paragraphs.index(h1)
    image_paragraph = paragraphs[h1_index + 1]

    for run in list(h1.runs):
        run._r.getparent().remove(run._r)
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.add_run(TITLE_PAGE_LINE_1)
    h1.add_run().add_break()
    h1.add_run(TITLE_PAGE_LINE_2)

    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.space_before = Inches(1.7)
    logo_shape = doc.inline_shapes[0]
    logo_shape.width = TITLE_LOGO_SIZE
    logo_shape.height = TITLE_LOGO_SIZE

    revision_p = doc.add_paragraph()
    revision_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    revision_p.paragraph_format.space_before = Inches(1.6)
    run = revision_p.add_run(TITLE_PAGE_REVISION)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    image_paragraph._p.addnext(revision_p._p)

    doc.save(str(docx_path))


_LEVEL_1_LIST_RE = re.compile(r'<w:lvl w:ilvl="1">.*?</w:lvl>', re.DOTALL)


def use_letters_for_nested_lists(docx_path: Path) -> None:
    """Render the 2nd level of every ordered list ("a. b. c.") instead of decimal.

    The Markdown source encodes nested list items as real nested ordered
    lists (`   1.`, `   2.`, ...) — CommonMark/GFM has no letter-marker list
    syntax, so writing literal "a." "b." in the source doesn't produce real
    list items at all; Pandoc (and GitHub) just treat them as continuation
    text of the parent item with no extra indentation. The numbered-vs-
    lettered visual distinction the source markdown wants is instead applied
    here, in the docx's own numbering definitions, which Word supports
    natively via `w:numFmt`. Pandoc generates numbering.xml fresh for every
    build (it isn't inherited from reference.docx), so this has to run as a
    post-process step on each generated docx, not baked into the template.
    """
    tmp_path = docx_path.with_suffix(docx_path.suffix + ".tmp")
    with zipfile.ZipFile(docx_path, "r") as src, zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as dst:
        for item in src.infolist():
            data = src.read(item.filename)
            if item.filename == "word/numbering.xml":
                xml_text = data.decode("utf-8")
                xml_text = _LEVEL_1_LIST_RE.sub(
                    lambda m: m.group(0).replace('w:val="decimal"', 'w:val="lowerLetter"', 1),
                    xml_text,
                )
                data = xml_text.encode("utf-8")
            dst.writestr(item, data)
    shutil.move(str(tmp_path), str(docx_path))
