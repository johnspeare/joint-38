#!/usr/bin/env python3
"""Regenerate pipeline/reference.docx — the Pandoc styling template.

Run this whenever the brand look (heading colors, footer text, logo) needs
to change. The output is a checked-in binary (reference.docx); this script
is how you reproduce or edit it instead of hand-editing XML.

Requires: pandoc on PATH, `pip install python-docx`.
"""

from __future__ import annotations

import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

PIPELINE_DIR = Path(__file__).resolve().parent
REPO_ROOT = PIPELINE_DIR.parent
ASSETS_DIR = REPO_ROOT / "FD-SOGs-assets"
DEFAULT_REFERENCE = PIPELINE_DIR / "_pandoc-default-reference.docx"
OUTPUT = PIPELINE_DIR / "reference.docx"

BRAND_RED = RGBColor(0x8B, 0x00, 0x00)
FOOTER_NOTICE = "STANDARD OPERATING GUIDELINES — UNCONTROLLED WHEN PRINTED"


def add_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def build_footer_paragraph(paragraph) -> None:
    paragraph.text = ""
    run = paragraph.add_run(FOOTER_NOTICE)
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    paragraph.add_run("\t\t")
    paragraph.add_run("Page ")
    add_field(paragraph, "PAGE")
    paragraph.add_run(" of ")
    add_field(paragraph, "NUMPAGES")
    for r in paragraph.runs[-4:]:
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_thin_table_borders(doc: Document, style_name: str = "Table") -> None:
    """Give the table style hairline borders on every edge, incl. between cells.

    Pandoc's default reference.docx defines a "Table" style with no borders
    at all, so every table renders as bare, unruled text columns. python-docx
    has no high-level API for table-*style* borders (only per-table/per-cell
    overrides), so this edits the style's <w:tblPr> XML directly.
    """
    style_element = doc.styles[style_name].element
    tbl_pr = style_element.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        style_element.append(tbl_pr)

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_el = OxmlElement(f"w:{edge}")
        edge_el.set(qn("w:val"), "single")
        edge_el.set(qn("w:sz"), "4")  # 4 eighths-of-a-point = 0.5pt hairline
        edge_el.set(qn("w:space"), "0")
        edge_el.set(qn("w:color"), "999999")
        borders.append(edge_el)
    tbl_pr.append(borders)


def main() -> None:
    subprocess.run(
        ["pandoc", "-o", str(DEFAULT_REFERENCE), "--print-default-data-file", "reference.docx"],
        check=True,
    )

    doc = Document(str(DEFAULT_REFERENCE))

    # Pandoc's default reference.docx points body/heading text at "Aptos" via
    # the docx theme — a Microsoft-only font not present on either macOS or
    # Linux LibreOffice. Both substitute for it, but with *different*
    # fallback fonts, which changes line-wrapping and therefore pagination
    # between build machines (found via CI: 79 pages locally on macOS vs. 92
    # on the Ubuntu runner, same source, same LibreOffice version). Pin an
    # explicit font that LibreOffice bundles identically on every platform
    # instead of relying on theme/system substitution.
    BODY_FONT = "Liberation Sans"
    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3", "Heading 4", "Title"):
        if style_name in doc.styles:
            doc.styles[style_name].font.name = BODY_FONT

    heading_sizes = {"Heading 1": 24, "Heading 2": 16, "Heading 3": 13, "Heading 4": 11.5}
    for name, size in heading_sizes.items():
        style = doc.styles[name]
        style.font.color.rgb = BRAND_RED
        style.font.size = Pt(size)
        style.font.bold = True

    if "Title" in doc.styles:
        title_style = doc.styles["Title"]
        title_style.font.color.rgb = BRAND_RED
        title_style.font.size = Pt(28)

    doc.styles["Heading 1"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_thin_table_borders(doc)

    section = doc.sections[0]
    section.different_first_page_header_footer = True

    build_footer_paragraph(section.footer.paragraphs[0])
    section.footer.is_linked_to_previous = False
    build_footer_paragraph(section.first_page_footer.paragraphs[0])

    # Small running badge in the header, interior pages only — the title
    # page carries the large logo in the body content instead.
    header = section.header
    header.is_linked_to_previous = False
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header_p.add_run()
    header_run.add_picture(str(ASSETS_DIR / "38-logo.png"), height=Inches(0.45))
    section.first_page_header.paragraphs[0].text = ""

    doc.save(str(OUTPUT))
    DEFAULT_REFERENCE.unlink()
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    sys.exit(main())
